from rest_framework.views import APIView
from django.http import JsonResponse
from rest_framework.response import Response
from rest_framework import status
from .models import User, Course, CourseAndUser, UserProgress, Material, Lesson, Compound, Feedback
from django.contrib.auth import authenticate, login, logout
from rest_framework_simplejwt.tokens import RefreshToken
from django.core.serializers import serialize
from django.contrib.auth.hashers import make_password
import base64, json
from django.db.models import Q
import random, string
from django.core.mail import EmailMessage
import os
from django.http import FileResponse, Http404
from django.conf import settings
from urllib.parse import quote
from docx import Document

class Registration(APIView):
    def post(self, request):
        firstName = request.data.get('firstName')
        secondName = request.data.get('secondName')
        lastName = request.data.get('lastName')
        email = request.data.get('email')
        role=request.data.get('role')

        if User.objects.filter(email=email).exists():
            return Response({'error': 'Пользователь с таким email уже существует'}, status=status.HTTP_400_BAD_REQUEST)
        password = self.generatePassword()
        print(password)
        hashed_password = make_password(password)

        user = User.objects.create(
            firstName=firstName,
            secondName=secondName,
            lastName=lastName,
            email=email,
            password=hashed_password,
            role=role
        )
        if role == 0:
            email_message = EmailMessage(
                    'Регистрация студента',
                    f'Для авторизации вам понадобится ваш email и пароль - {password}, желаю потратить побольше денег на курсы искатель',
                    to=[email]
                )
        else:
            email_message = EmailMessage(
                'Регистрация учителя',
                f'Для авторизации вам понадобится ваш email и пароль - {password}, желаю приобрести побольше денег на родителях бедных детишек',
                to=[email]
            )
        
        email_message.send()

        return Response({'message': 'Регистрация успешна',},
                         status=status.HTTP_201_CREATED)
                         

    def generatePassword(self, length=12):
        lower = string.ascii_lowercase
        upper = string.ascii_uppercase
        digits = string.digits

        all_chars = lower + upper + digits 

        password = [
            random.choice(lower),
            random.choice(upper),
            random.choice(digits)
        ]

        password += random.choices(all_chars, k=length - 4)

        random.shuffle(password)

        return ''.join(password)
    
class Login(APIView):
    def post(self, request):
        email = request.data.get('email')
        password = request.data.get('password')

        user = authenticate(request, email=email, password=password)
        if user is None:
            return Response({'error': 'Неверное имя пользователя или пароль'}, status=status.HTTP_401_UNAUTHORIZED)

        refresh = RefreshToken.for_user(user)

        return Response({
            'message': 'Авторизация успешна',
            'tokens': {
                'refresh': str(refresh),
                'access': str(refresh.access_token),
            }
        }, status=status.HTTP_200_OK)
    
class Logout(APIView):
    def post(self, request):
        logout(request)
        return Response({'message': 'Выход выполнен успешно'}, status=status.HTTP_200_OK)
    
class getUsersByCourse(APIView):
    def get(self, request):
        courseId = request.query_params.get('courseId')

        course = Course.objects.get(id = courseId)

        existing_user_ids = CourseAndUser.objects.filter(course=course).values_list('user_id', flat=True)
        usrs = User.objects.exclude(id__in=existing_user_ids).filter(role = 0)
        users = []
    
        for user in usrs:
            users.append({
                'id': user.id,
                'FIO': f'{user.secondName} {user.firstName} {user.lastName}',
                'email': user.email
            })
        print(users)
        return Response({'users': users}, status=status.HTTP_200_OK)
    
class getUser(APIView):
    def post(self, request):
        userId = request.data.get('userId')

        user = User.objects.get(id = userId)

        return Response({'role': user.role}, status=status.HTTP_200_OK)
    
#Добавить пользователя на курс
class inviteUserOnCourse(APIView):
    def post(self, request):
        userId = request.data.get('userId')
        courseId = request.data.get('courseId')
        user = User.objects.get(id = userId)
        course = Course.objects.get(id = courseId)

        CourseAndUser.objects.create(user = user, course = course)

        lessons = Lesson.objects.filter(course = course)

        for lesson in lessons:
            materials = Material.objects.filter(lesson = lesson)
            for material in materials:
                UserProgress.objects.create(user = user, course = course, lesson = lesson, material = material)
        
        return Response(status=status.HTTP_201_CREATED)

#Создать урок
class createTask(APIView):
    def post(self, request):
        courseId = request.data.get('courseId')
        name = request.data.get('name')

        course = Course.objects.get(id = courseId)
        if Lesson.objects.filter(course=course, name=name).exists():
            return Response({'message': 'Урок с таким названием уже есть'}, status=status.HTTP_400_BAD_REQUEST)
        lesson = Lesson.objects.create(name = name, course = course)

        return Response({'lessonId': lesson.id},status=status.HTTP_201_CREATED)

#Создать материал
class createMaterial(APIView):
    def post(self, request):
        courseId = request.data.get('courseId')
        lessonName = request.data.get('lessonName')
        type = request.data.get('type')
        name = request.data.get('name')
        file = request.FILES.get('file')

        course = Course.objects.get(id = courseId)

        lesson = Lesson.objects.get(course = course, name = lessonName)

        if Material.objects.filter(lesson = lesson, name=name).exists():
            return Response({'message': 'Материал с таким названием уже есть'}, status=status.HTTP_400_BAD_REQUEST)

        if type == '4':
            link = request.data.get('link')
            material = Material.objects.create(name = name, type = type, link = link, lesson = lesson)

            return Response(status=status.HTTP_201_CREATED)

        ext = os.path.splitext(file.name)[1]
        fileName = f'{name}{ext}'

        lessonName = lesson.name
        courseName = lesson.course.name

        target_directory = os.path.join(os.getcwd(), f'../files/courses/{courseName}/{lessonName}')
        os.makedirs(target_directory, exist_ok=True)

        target_path = os.path.join(target_directory, fileName)
        
        with open(target_path, 'wb') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        material = Material.objects.create(name = name, type = type,file = f'../files/courses/{courseName}/{lessonName}/{fileName}', lesson = lesson)
        if type != 1:
            usersOnCourse = CourseAndUser.objects.filter(course = course)
            for userOnCourse in usersOnCourse:
                UserProgress.objects.create(user = userOnCourse.user, course = course, lesson = lesson, material = material)

        tmp = {
            'name': name,
            'type': type,
            'id': material.id
        }
        return Response({'material': tmp},status=status.HTTP_201_CREATED)
    
#Создать курс
class createCourse(APIView):
    def post(self, request):
        teacherId = request.data.get('userId')
        name = request.data.get('name')
        description = request.data.get('description')
        places = request.data.get('places')
        price = request.data.get('price')
        salePrice = request.data.get('salePrice')
        sale = request.data.get('sale')
        credit = request.data.get('credit')
        picture = request.FILES.get('picture')
        compounds = json.loads(request.data.get('compounds'))
        teacher = User.objects.get(id = teacherId)

        if Course.objects.filter(name=name).exists():
                return Response(
                    {'error': 'Курс с таким названием уже существует'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        ext = os.path.splitext(picture.name)[1]
        picName = f'{name}{ext}'

        target_directory = os.path.join(os.getcwd(), f'../nuxt-app/public/images/courses')
        os.makedirs(target_directory, exist_ok=True)

        target_path = os.path.join(target_directory, picName)
        
        with open(target_path, 'wb') as destination:
            for chunk in picture.chunks():
                destination.write(chunk)

        course = Course.objects.create(teacher = teacher, name = name, description = description, places = places, price = price,
                              salePrice = salePrice, sale = sale, credit = credit, picture = f'/images/courses/{picName}')
        
        for compound in compounds:
            if (compound.strip() !=''):
                Compound.objects.create(course = course, name = compound)

        return Response(status=status.HTTP_201_CREATED)


#Курсы студента
class getCourseByUser(APIView):
    def post(self, request):
        userId = request.data.get('userId')
        user = User.objects.get(id = userId)

        courses = []
        crs = CourseAndUser.objects.filter(user = user)

        for courseAndUser in crs:
            course = courseAndUser.course
            check = UserProgress.objects.filter(user=user,course=course).exists()
            if check:
                userProgress = UserProgress.objects.filter(user = user, course = course)
                done = userProgress.filter(done = True)
                progress = len(done) / len(userProgress) * 100
            else:
                progress = 0
            courses.append({
                'id': course.id,
                'picture': course.picture,
                'name': course.name,
                'progress': progress
            })
        
        return Response({'courses': courses}, status=status.HTTP_200_OK)

#Курс студента
class getCourseForUser(APIView):
    def get(self, request):
        userId = request.query_params.get('userId')
        courseId = request.query_params.get('courseId')
        user = User.objects.get(id = userId)
        crs = Course.objects.get(id = courseId)
        
        course = {}
        lessons = Lesson.objects.filter(course = crs)
        for lesson in lessons:
            mtrls = Material.objects.filter(lesson = lesson)
            materials = []
            for material in mtrls:
                if material.type == '4':
                    materials.append({
                        'id': material.id,
                        'name': material.name,
                        'type': material.type,
                        'link': material.link
                    })
                else:
                    materials.append({
                    'id': material.id,
                    'name': material.name,
                    'type': material.type,
                    'file': material.file
                })
            course[lesson.name] = materials


        return Response({'course': course}, status=status.HTTP_200_OK)

#Скачать файл
class downloadFile(APIView):
    def post(self, request):
        file_path = request.data.get('path')

        if not os.path.exists(file_path):
            raise Http404("Файл не найден")

        # Получаем имя файла из пути
        file_name = os.path.basename(file_path)

        # Открываем файл
        file = open(file_path, 'rb')

        # Создаём ответ с файлом
        response = FileResponse(file)

        # Устанавливаем Content-Disposition для скачивания с оригинальным названием
        response['Content-Disposition'] = f'attachment; filename="{quote(file_name)}'

        return response
        
#Прикрепить файл
class uploadFile(APIView):
    def post(self, request):
        userId = request.data.get('userId')
        materialId = request.data.get('materialId')
        file = request.FILES.get('file')

        user = User.objects.get(id = userId)
        material = Material.objects.get(id = materialId)
        lesson = material.lesson
        lessonName = lesson.name
        courseName = lesson.course.name

        target_directory = os.path.join(os.getcwd(), f'../files/users/{userId}/{courseName}/{lessonName}')
        os.makedirs(target_directory, exist_ok=True)

        target_path = os.path.join(target_directory, file.name)
        
        with open(target_path, 'wb') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        
        userProgress = UserProgress.objects.get(user = user, lesson = lesson, material = material)
        userProgress.file = f'../files/users/{userId}/{courseName}/{lessonName}/{file.name}'
        userProgress.needToCheck = True
        userProgress.save()

        return Response(status=status.HTTP_200_OK)
    
#Домашка для проверки
class getFilesForTeacher(APIView):
    def post(self, request):
        courseId = request.data.get('courseId')
        course = Course.objects.get(id = courseId)

        userProgress = UserProgress.objects.filter(course = course, needToCheck = True)
        homeworks = []

        for progress in userProgress:
            user = progress.user
            fio = f'{user.secondName} {user.firstName} {user.lastName}'
            lesson = progress.lesson
            material = progress.material
            file = progress.file
            homeworks.append({
                'id': progress.id,
                'fio': fio,
                'lessonName': lesson.name,
                'materialName': material.name,
                'file': file
            })
        
        return Response({'homeworks': homeworks}, status=status.HTTP_200_OK)
    
#Кнопка проверки
class checked(APIView):
    def post(self, request):
        userProgressId = request.data.get('userProgressId')

        userProgress = UserProgress.objects.get(id = userProgressId)

        userProgress.needToCheck = False
        userProgress.done = True
        userProgress.save()

        return Response(status=status.HTTP_200_OK)

#Курсы учителя
class getCoursesForTeacher(APIView):
    def get(self, request):
        teacherId = request.query_params.get('teacherId')
        teacher = User.objects.get(id = teacherId)

        crs = Course.objects.filter(teacher = teacher)

        courses = []

        for course in crs:
            needToCheck = UserProgress.objects.filter(course = course, needToCheck = True).count()
            courses.append({
                'id': course.id,
                'name': course.name,
                'needToCheck': needToCheck,
                'picture': course.picture
            })
        
        return Response({'courses': courses}, status=status.HTTP_200_OK)


#Учебные материалы для учителя
class getCourseForTeacher(APIView):
    def post(self, request):
        courseId = request.data.get('courseId')
        crs = Course.objects.get(id = courseId)

        course = {}
        lessons = Lesson.objects.filter(course = crs)
        for lesson in lessons:
            mtrls = Material.objects.filter(lesson = lesson)
            materials = []
            for material in mtrls:
                materials.append({
                    'id': material.id,
                    'name': material.name,
                    'type': material.type,
                    
                })
               
            course[lesson.id] = {'id': lesson.id, 'name': lesson.name, 'materials':materials}

        return Response({'course': course, 'courseName': crs.name}, status=status.HTTP_200_OK)
    
class updateLesson(APIView):
    def post(self, request):
        courseId = request.data.get('courseId')
        lessonName = request.data.get('oldName')
        newName = request.data.get('newName')
        print(courseId)
        course = Course.objects.get(id = courseId)

        lesson = Lesson.objects.get(course = course,name = lessonName)
        lesson.name = newName

        lesson.save()

        return Response(status=status.HTTP_200_OK)
    
class updateMaterial(APIView):
    def post(self, request):
        materialId = request.data.get('materialId')
        newName = request.data.get('newName')

        material = Material.objects.get(id = materialId)
        material.name = newName

        material.save()

        return Response(status=status.HTTP_200_OK)


class getDoneUserProgress(APIView):
    def post(self, request):
        userId = request.data.get('userId')

        user = User.objects.get(id = userId)

        userProgress = UserProgress.objects.filter(user = user, done = True)

        materials = [progress.material.id for progress in userProgress]

        return Response({'materials': materials}, status=status.HTTP_200_OK)
    
class getCourseForMain(APIView):
    def get(self, request):
        crss = Course.objects.all()

        courses = []
        for crs in crss:
            compounds=[]
            cmpnds = Compound.objects.filter(course=crs)
            for compound in cmpnds:
                compounds.append({
                    'id':compound.id,
                    'name':compound.name
                })
            courses.append(
                {
                    'name': crs.name,
                    'description': crs.description,
                    'picture': crs.picture,
                    'price': crs.price,
                    'salePrice': crs.salePrice,
                    'credit': crs.credit,
                    'places': crs.places,
                    'sale': crs.sale,
                    'compounds': compounds
                }
            )   

        return Response(courses,status=status.HTTP_200_OK)
    
class sendEmailToAdmin(APIView):
    def post(self, request):
        fio = request.data.get('fio')
        number = request.data.get('number')

        email_message = EmailMessage(
                    'Анкета',
                    f'Подали заявку телефон - {number}, ФИО - {fio}',
                    to=['vladi.popova2004@yandex.ru'])
        
        email_message.send()
        
        return Response(status=status.HTTP_200_OK)
    
class deleteTeacher(APIView):
    def post(self, request):
        id = request.data.get('id')

        user = User.objects.get(id=id)

        user.delete()

        return Response(status=status.HTTP_200_OK)
    
class getTeachers(APIView):
    def get(self, request):
        tchrs = User.objects.filter(role=1)
        teachers = []

        for teacher in tchrs:
            teachers.append(
                {
                    'id': teacher.id,
                    'FIO': f'{teacher.secondName} {teacher.firstName} {teacher.lastName}',
                    'email': teacher.email
                }
            )
        return Response(teachers, status=status.HTTP_200_OK)
    
class getFeedback(APIView):
    def get(self, request):
        fds = Feedback.objects.filter(visible=True)
        feedbacks = []
        for feedback in fds:
            feedbacks.append(
                {
                    'id': feedback.id,
                    'sender': feedback.sender,
                    'text': feedback.text
                }
            )
        
        return Response(feedbacks, status=status.HTTP_200_OK)

class createComment(APIView):
    def post(self, request):
        sender = request.data.get('sender')
        text = request.data.get('text')

        all = Feedback.objects.all().count()
        if all >= 3:
            third = Feedback.objects.order_by('-id')[2:3].first()
            third.visible = False
            third.save()
        Feedback.objects.create(sender=sender, text=text)

        return Response(status=status.HTTP_200_OK)
    
class parseFile(APIView):
    def post(self, request):
        matId = request.data.get('matId')
        material = Material.objects.get(id=matId)
        materialName = material.name
        doc = Document(material.file)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        image_folder = f'../nuxt-app/public/images/tests/{materialName}'
        if not os.path.exists(image_folder):
            os.makedirs(image_folder)

        image_relations = []
        
        for rel in doc.part._rels.values():
            if "image" in rel.target_ref:
                image_relations.append(rel)

        saved_images = []

        test_data = []
        i = 0
        image_index = 0  

        while i < len(lines):
            line = lines[i]

            if line.startswith("Вопрос"):
                question = {
                    "question": "",
                    "type": "",
                    "image": "",
                    "answers": []
                }

                question["question"] = line.split(":", 1)[1].strip('«»"\' ')
                i += 1

                if i < len(lines) and lines[i].startswith("Тип вопроса"):
                    q_type = lines[i].split(":", 1)[1].strip().lower()
                    if 'открытый' in q_type:
                        question["type"] = "open"
                    elif 'один' in q_type or 'single' in q_type:
                        question["type"] = "single"
                    elif 'несколько' in q_type or 'multiple' in q_type:
                        question["type"] = "multiple"
                    else:
                        question["type"] = "unknown"
                    i += 1
                else:
                    question["type"] = "unknown"

                if i < len(lines) and lines[i].startswith("Изображение"):
                    img_flag = lines[i].split(":", 1)[1].strip().strip('«»"\' ')
                    print(i,img_flag)
                    if img_flag == "+" and image_index < len(image_relations):
                        rel = image_relations[image_index]
                        image_blob = rel.target_part.blob
                        ext = rel.target_part.content_type.split("/")[-1]
                        img_filename = f"{image_folder}/img_{image_index}.{ext}"
                        with open(img_filename, "wb") as f:
                            f.write(image_blob)
                        question["image"] = f'/images/tests/{materialName}/img_{image_index}.{ext}'
                        saved_images.append(img_filename)
                        image_index += 1
                    elif img_flag == "-" or img_flag == "":
                        question["image"] = "-"
                    else:
                        question["image"] = "Ошибка: не найдено изображение"
                    i += 1

                if i < len(lines) and lines[i] == "Ответы:":
                    i += 1
                    while i < len(lines) and lines[i].startswith(("+", "-")):
                        answer_text = lines[i][2:].strip()
                        is_correct = lines[i][0] == "+"
                        question["answers"].append({
                            "text": answer_text,
                            "correct": is_correct
                        })
                        i += 1
                else:
                    i += 1

                test_data.append(question)
            else:
                i += 1
        return Response(test_data,status=status.HTTP_200_OK)
    
class welcomeToTeacher(APIView):
    def post(self, request):
        userId = request.data.get('userId')
        matId = request.data.get('matId')
        answers = request.data.get('test')
        questions = request.data.get('questions')

        user = User.objects.get(id = userId)
        userName = f'{user.secondName} {user.firstName[0]}.{user.lastName[0]}.'

        material = Material.objects.get(id=matId)
        materialName = material.name

        filePath = f'../files/courses/{material.lesson.course.name}/{material.lesson.name}/тесты/'

        if not os.path.exists(filePath):
            os.makedirs(filePath)

        doc = Document()

        doc.add_heading('Тест', level=1)
        doc.add_heading(f'«{materialName}»', level=2)

        p = doc.add_paragraph()
        p.add_run(f'Студент « {userName} »').italic = True
        for i, (question, answer) in enumerate(zip(questions, answers), start=1):
            doc.add_paragraph(f'Вопрос {i}: {question}')
            doc.add_paragraph('Ответ:')

            if isinstance(answer, list):
                for option in answer:
                    marker = '+' if option['correct'] else '-'
                    doc.add_paragraph(f' {marker}  {option["text"]}')
            elif isinstance(answer, dict):
                marker = '+' if answer['correct'] else '-'
                doc.add_paragraph(f' {marker}  {answer["text"]}')

        doc.save(f'{filePath}{userName}_{materialName}.docx')

        userProgress = UserProgress.objects.get(user = user, lesson = material.lesson, material = material)
        userProgress.file = f'{filePath}{userName}_{materialName}.docx'
        userProgress.needToCheck = True
        userProgress.save()
        
        return Response(status=status.HTTP_200_OK)

class deleteMaterial(APIView):
    def post(self, request):
        matId = request.data.get('matId')
        material = Material.objects.get(id=matId)
        material.delete()
        return Response(status=status.HTTP_200_OK)

class deleteLesson(APIView):
    def post(self, request):
        lesId = request.data.get('lesId')
        courseId = request.data.get('courseId')
        course = Course.objects.get(id=courseId)
        lesson = Lesson.objects.get(id=lesId, course=course)
        lesson.delete()
        return Response(status=status.HTTP_200_OK)

class deleteCourse(APIView):
    def post(self, request):
        courseId = request.data.get('courseId')
        course = Course.objects.get(id=courseId)
        course.delete()
        return Response(status=status.HTTP_200_OK)


        



        



                






